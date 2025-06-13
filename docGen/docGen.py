# docGen.py
# probably the best vibe code yet
# pandas must be installed
# python-docx must be installed
# reportlab must be installed
# python -m tkinter
#___________________________________________________________________________________

import os
import pandas as pd
import tkinter as tk
from tkinter import filedialog
from datetime import datetime
from docx import Document
from reportlab.lib.pagesizes import letter, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, PageBreak
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet

# Define Tabloid size manually (11x17 inches)
TABLOID = (11 * 72, 17 * 72)  # 1 inch = 72 points

# Page size options
PAGE_SIZES = {"Letter": letter, "Tabloid": TABLOID}

# Character limit for detecting oversized columns
CHARACTER_LIMIT = 120

# Max characters per line before inserting a newline
MAX_CHAR_PER_LINE = 60

def select_file():
    """Opens a file dialog for the user to select a .csv or .txt file."""
    root = tk.Tk()
    root.withdraw()
    file_path = filedialog.askopenfilename(filetypes=[("CSV & TXT Files", "*.csv;*.txt")])
    return file_path

def parse_csv(file_path):
    """Parses a CSV file into multiple tables based on blank columns and rows."""
    df = pd.read_csv(file_path, header=None)
    tables = []
    current_table = []
    
    for _, row in df.iterrows():
        if row.isnull().all():  # Blank row indicates end of a table
            if current_table:
                tables.append(pd.DataFrame(current_table))
                current_table = []
        else:
            current_table.append(row.dropna().tolist())  # Drop blank columns
    
    if current_table:  # Add last table if not empty
        tables.append(pd.DataFrame(current_table))
    
    return tables

def parse_txt(file_path):
    """Reads a TXT file and returns its content as a list of lines."""
    with open(file_path, "r", encoding="utf-8") as file:
        return file.readlines()

def create_title_page(doc, file_name, tables):
    """Creates a title page with document name, date, and table of contents."""
    doc.add_paragraph(f"Document: {file_name}")
    doc.add_paragraph(f"Date: {datetime.today().strftime('%Y-%m-%d')}")
    doc.add_paragraph("\nTable of Contents:")
    
    for i, _ in enumerate(tables, start=1):
        doc.add_paragraph(f"Table {i}")

    doc.add_page_break()

def split_text_into_lines(text, max_chars_per_line=MAX_CHAR_PER_LINE):
    """Splits long text into multiple lines at natural breakpoints."""
    words = text.split()
    lines = []
    current_line = ""

    for word in words:
        if len(current_line) + len(word) + 1 > max_chars_per_line:
            lines.append(current_line.strip())
            current_line = word
        else:
            current_line += " " + word

    if current_line:
        lines.append(current_line.strip())

    return "\n".join(lines)

def split_table_by_content(table):
    """Splits tables based on columns with oversized content."""
    oversized_columns = [col for col in table.columns if table[col].astype(str).apply(len).max() > CHARACTER_LIMIT]
    
    if not oversized_columns:
        return [table]  # No need to split
    
    normal_columns = [col for col in table.columns if col not in oversized_columns]
    
    split_tables = []
    if normal_columns:
        split_tables.append(table[normal_columns])  # Table with normal columns
    
    for col in oversized_columns:
        table[col] = table[col].astype(str).apply(split_text_into_lines)  # Apply text splitting
        split_tables.append(table[[col]])  # Each oversized column gets its own page
    
    return split_tables

def create_pdf(file_name, page_size, orientation, tables):
    """Generates a PDF with formatted tables and title page."""
    pdf_file = file_name.replace(".csv", ".pdf").replace(".txt", ".pdf")
    doc = SimpleDocTemplate(pdf_file, pagesize=landscape(PAGE_SIZES[page_size]) if orientation == "Landscape" else PAGE_SIZES[page_size])
    elements = []
    styles = getSampleStyleSheet()
    
    # Title Page
    elements.append(Paragraph(f"Document: {file_name}", styles["Title"]))
    elements.append(Paragraph(f"Date: {datetime.today().strftime('%Y-%m-%d')}", styles["Normal"]))
    elements.append(Paragraph("Table of Contents:", styles["Heading2"]))
    
    for i, _ in enumerate(tables, start=1):
        elements.append(Paragraph(f"Table {i}", styles["Normal"]))
    
    elements.append(PageBreak())

    # Tables
    for i, table in enumerate(tables, start=1):
        elements.append(Paragraph(f"Table {i}", styles["Heading2"]))
        
        # Split tables based on oversized content
        split_tables = split_table_by_content(table)
        
        for sub_table in split_tables:
            data = [sub_table.columns.tolist()] + sub_table.values.tolist()
            table_obj = Table(data)
            table_obj.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('WORDWRAP', (0, 0), (-1, -1)),  # Ensure text wraps instead of cutting off
                ('LEADING', (0, 0), (-1, -1), 18)  # Increase row height for oversized columns (3 lines per row)
            ]))
            elements.append(table_obj)
            elements.append(PageBreak())

    doc.build(elements)
    print(f"PDF saved as {pdf_file}")

def create_docx(file_name, tables):
    """Generates a DOCX with formatted tables and title page."""
    docx_file = file_name.replace(".csv", ".docx").replace(".txt", ".docx")
    doc = Document()
    
    # Title Page
    create_title_page(doc, file_name, tables)

    # Tables
    for i, table in enumerate(tables, start=1):
        doc.add_paragraph(f"Table {i}")
        
        # Split tables based on oversized content
        split_tables = split_table_by_content(table)
        
        for sub_table in split_tables:
            doc_table = doc.add_table(rows=len(sub_table) + 1, cols=len(sub_table.columns))
            
            # Add headers
            for j, header in enumerate(sub_table.columns):
                doc_table.cell(0, j).text = str(header)
            
            # Add rows
            for row_idx, row in enumerate(sub_table.values, start=1):
                for col_idx, value in enumerate(row):
                    doc_table.cell(row_idx, col_idx).text = str(value)
            
            doc.add_page_break()

    doc.save(docx_file)
    print(f"DOCX saved as {docx_file}")

def main():
    """Main function to run the program."""
    file_path = select_file()
    if not file_path:
        print("No file selected. Exiting.")
        return
    
    file_name = os.path.basename(file_path)
    file_ext = os.path.splitext(file_name)[1].lower()
    
    output_format = input("Choose output format (PDF/DOCX): ").strip().upper()
    page_size = input("Choose page size (Letter/Tabloid): ").strip().title()
    orientation = input("Choose orientation (Portrait/Landscape): ").strip().title()
    
    tables = parse_csv(file_path) if file_ext == ".csv" else [parse_txt(file_path)]
    
    if output_format == "PDF":
        create_pdf(file_name, page_size, orientation, tables)
    elif output_format == "DOCX":
        create_docx(file_name, tables)

if __name__ == "__main__":
    main()
